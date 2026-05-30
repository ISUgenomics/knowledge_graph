#!/usr/bin/env python3
"""
extract_text.py — Generalized file text extraction.

Handles PDF, DOCX, XLSX, CSV, HTML, TXT, and MD files.
Uses the best available tool for each format (best practice #13).
Returns structured output with format, char count, and error info.
"""

import csv
import io
import re
from pathlib import Path


def extract_text(path: str, max_chars: int = 50000) -> dict:
    """
    Extract plain text from a file.

    Returns:
        {
            "text": str,
            "format": str,
            "chars": int,
            "pages": int | None,
            "error": str | None,
        }
    """
    p = Path(path)
    if not p.exists():
        return {"text": "", "format": "unknown", "chars": 0, "pages": None,
                "error": f"File not found: {path}"}

    suffix = p.suffix.lower()
    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".xlsx": _extract_xlsx,
        ".csv": _extract_csv,
        ".html": _extract_html,
        ".htm": _extract_html,
        ".txt": _extract_plain,
        ".md": _extract_plain,
        ".tsv": _extract_csv,
    }

    extractor = extractors.get(suffix)
    if not extractor:
        return {"text": "", "format": suffix.lstrip(".") or "unknown", "chars": 0,
                "pages": None, "error": f"Unsupported format: {suffix}"}

    try:
        result = extractor(str(p), max_chars)
        if not result.get("format"):
            result["format"] = suffix.lstrip(".")
        return result
    except Exception as e:
        return {"text": "", "format": suffix.lstrip("."), "chars": 0,
                "pages": None, "error": str(e)}


def _extract_pdf(path: str, max_chars: int) -> dict:
    """Extract text from PDF using PyMuPDF or pdftotext."""
    pdf_bytes = Path(path).read_bytes()
    if len(pdf_bytes) < 100 or pdf_bytes[:5] != b"%PDF-":
        return {"text": "", "format": "pdf", "chars": 0, "pages": None,
                "error": "Not a valid PDF file"}

    # Try PyMuPDF first
    try:
        import fitz
        doc = fitz.open(path)
        pages = len(doc)
        text = "\n".join(page.get_text() for page in doc).strip()
        doc.close()
        if text:
            return {"text": text[:max_chars], "format": "pdf", "chars": len(text),
                    "pages": pages, "error": None}
    except ImportError:
        pass
    except Exception:
        pass

    # Fallback: pdftotext CLI
    try:
        import subprocess
        proc = subprocess.run(["pdftotext", path, "-"],
                              capture_output=True, timeout=30)
        if proc.returncode == 0:
            text = proc.stdout.decode("utf-8", errors="replace").strip()
            if text:
                return {"text": text[:max_chars], "format": "pdf",
                        "chars": len(text), "pages": None, "error": None}
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    return {"text": "", "format": "pdf", "chars": 0, "pages": None,
            "error": "No PDF extractor available (install pymupdf or pdftotext)"}


def _extract_docx(path: str, max_chars: int) -> dict:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        text = "\n".join(paragraphs)
        return {"text": text[:max_chars], "format": "docx",
                "chars": len(text), "pages": None, "error": None}
    except ImportError:
        pass

    # Fallback: pandoc CLI
    try:
        import subprocess
        proc = subprocess.run(["pandoc", "-f", "docx", "-t", "plain", path],
                              capture_output=True, text=True, timeout=30)
        if proc.returncode == 0 and proc.stdout.strip():
            text = proc.stdout.strip()
            return {"text": text[:max_chars], "format": "docx",
                    "chars": len(text), "pages": None, "error": None}
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    return {"text": "", "format": "docx", "chars": 0, "pages": None,
            "error": "No DOCX extractor available (install python-docx or pandoc)"}


def _extract_xlsx(path: str, max_chars: int) -> dict:
    """Extract text from XLSX as markdown tables."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(path, read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    rows.append(" | ".join(cells))
            if rows:
                parts.append(f"## Sheet: {sheet_name}\n")
                # Header + separator
                if len(rows) > 1:
                    parts.append(rows[0])
                    col_count = rows[0].count("|") + 1
                    parts.append(" | ".join(["---"] * col_count))
                    parts.extend(rows[1:])
                else:
                    parts.extend(rows)
                parts.append("")
        wb.close()
        text = "\n".join(parts)
        return {"text": text[:max_chars], "format": "xlsx",
                "chars": len(text), "pages": None, "error": None}
    except ImportError:
        return {"text": "", "format": "xlsx", "chars": 0, "pages": None,
                "error": "No XLSX extractor available (install openpyxl)"}


def _extract_csv(path: str, max_chars: int) -> dict:
    """Extract text from CSV/TSV as markdown table."""
    p = Path(path)
    delimiter = "\t" if p.suffix.lower() == ".tsv" else ","
    text_content = p.read_text(errors="replace")

    reader = csv.reader(io.StringIO(text_content), delimiter=delimiter)
    rows = []
    for row in reader:
        cells = [c.strip() for c in row]
        if any(cells):
            rows.append(" | ".join(cells))

    parts = []
    if rows:
        parts.append(rows[0])
        col_count = rows[0].count("|") + 1
        parts.append(" | ".join(["---"] * col_count))
        parts.extend(rows[1:])

    text = "\n".join(parts)
    return {"text": text[:max_chars], "format": p.suffix.lstrip("."),
            "chars": len(text), "pages": None, "error": None}


def _extract_html(path: str, max_chars: int) -> dict:
    """Extract text from HTML by stripping tags."""
    html = Path(path).read_text(errors="replace")
    text = _clean_html(html, max_chars)
    return {"text": text, "format": "html", "chars": len(text),
            "pages": None, "error": None}


def _extract_plain(path: str, max_chars: int) -> dict:
    """Read plain text file."""
    text = Path(path).read_text(errors="replace").strip()
    fmt = Path(path).suffix.lstrip(".") or "txt"
    return {"text": text[:max_chars], "format": fmt,
            "chars": len(text), "pages": None, "error": None}


def _clean_html(html: str, max_chars: int = 50000) -> str:
    """Strip HTML tags, scripts, styles and return clean text."""
    for tag in ("script", "style", "nav", "footer", "header", "noscript"):
        html = re.sub(rf"<{tag}[^>]*>.*?</{tag}>", "", html, flags=re.DOTALL | re.I)
    html = re.sub(r"<!--.*?-->", "", html, flags=re.DOTALL)
    html = re.sub(r"<(br|hr|/p|/div|/li|/tr|/h[1-6])[^>]*>", "\n", html, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", html)
    for ent, ch in [("&amp;", "&"), ("&lt;", "<"), ("&gt;", ">"),
                     ("&quot;", '"'), ("&#39;", "'"), ("&nbsp;", " ")]:
        text = text.replace(ent, ch)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    return text.strip()[:max_chars]


# ---------------------------------------------------------------------------
# Supported formats query
# ---------------------------------------------------------------------------
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".csv", ".tsv",
                        ".html", ".htm", ".txt", ".md"}


def is_supported(path: str) -> bool:
    """Check if a file extension is supported for text extraction."""
    return Path(path).suffix.lower() in SUPPORTED_EXTENSIONS


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import argparse
    import json
    parser = argparse.ArgumentParser(description="Extract text from a file")
    parser.add_argument("path", help="Path to file")
    args = parser.parse_args()
    result = extract_text(args.path)
    print(json.dumps({k: v for k, v in result.items() if k != "text"}, indent=2))
    if result["text"]:
        print(f"\n--- Text ({result['chars']} chars) ---")
        print(result["text"][:2000])
